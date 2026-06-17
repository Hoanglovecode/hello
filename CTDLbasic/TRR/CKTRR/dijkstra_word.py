import os
import sys
import re
import subprocess

# Tự động cài đặt thư viện python-docx nếu chưa có
try:
    import docx
except ImportError:
    print("[Aris Info] Thư viện 'python-docx' chưa được cài đặt. Đang tiến hành cài đặt tự động...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
        print("[Aris Info] Cài đặt thư viện 'python-docx' thành công!\n")
    except Exception as e:
        print(f"[Aris Error] Không thể tự động cài đặt python-docx: {e}")
        print("Sensei vui lòng cài đặt thủ công bằng lệnh: pip install python-docx")
        sys.exit(1)

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, color_hex):
    """Thiết lập màu nền cho ô trong Word (định dạng Hex, ví dụ: 'FF0000')"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_borders(cell, color="000000", sz="4", val="single"):
    """Tạo viền đen mảnh cho ô"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)  # sz tính bằng 1/8 pt, "4" = 0.5 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Thiết lập khoảng cách đệm (padding) trong ô (đơn vị dxa, 1/20 pt)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def parse_value(s):
    """Giải mã giá trị khoảng cách nhập vào, xử lý các dạng ký hiệu vô cực."""
    s = s.strip().lower()
    if s in ('inf', 'infinity', '∞', '-', 'none', 'null', ''):
        return float('inf')
    try:
        val = float(s)
        if val.is_integer():
            return int(val)
        return val
    except ValueError:
        return float('inf')

def format_T(T_set, all_vertices):
    """Định nghĩa cách viết tập hợp T (đỉnh chưa duyệt) tối ưu theo chuẩn của ảnh."""
    if not T_set:
        return "{}"
    T_sorted = [v for v in all_vertices if v in T_set]
    
    # Tìm các phân đoạn liên tục của chỉ số
    segments = []
    current_segment = []
    for v in T_sorted:
        idx = all_vertices.index(v)
        if not current_segment:
            current_segment.append(idx)
        else:
            if idx == current_segment[-1] + 1:
                current_segment.append(idx)
            else:
                segments.append(current_segment)
                current_segment = [idx]
    if current_segment:
        segments.append(current_segment)
        
    formatted_parts = []
    num_segments = len(segments)
    
    for seg in segments:
        seg_vertices = [all_vertices[i] for i in seg]
        length = len(seg)
        if length == 1:
            formatted_parts.append(seg_vertices[0])
        elif length == 2:
            formatted_parts.append(f"{seg_vertices[0]},{seg_vertices[1]}")
        elif length == 3:
            if num_segments == 1:
                formatted_parts.append(",".join(seg_vertices))
            else:
                formatted_parts.append(f"{seg_vertices[0]}..{seg_vertices[-1]}")
        else:
            formatted_parts.append(f"{seg_vertices[0]}...{seg_vertices[-1]}")
            
    return "{" + ",".join(formatted_parts) + "}"

def run_dijkstra(vertices, adjacency_matrix, start, end):
    """Chạy thuật toán Dijkstra và lưu lại vết từng bước lặp."""
    distances = {v: float('inf') for v in vertices}
    distances[start] = 0.0
    predecessors = {v: None for v in vertices}
    
    T = set(vertices)
    steps = []
    
    # Bước khởi tạo (Dòng 1 trong bảng)
    steps.append({
        'T': set(T),
        'selected_v': None,
        'distances': dict(distances),
        'highlight_v': None
    })
    
    while T:
        min_v = None
        min_dist = float('inf')
        for v in T:
            if distances[v] < min_dist:
                min_dist = distances[v]
                min_v = v
                
        if min_v is None or min_dist == float('inf'):
            break
            
        u = min_v
        T.remove(u)
        
        # Cập nhật khoảng cách các lân cận của u mà chưa duyệt (còn trong T)
        for v in T:
            weight = adjacency_matrix[u].get(v, float('inf'))
            if weight != float('inf'):
                new_dist = distances[u] + weight
                if new_dist < distances[v]:
                    distances[v] = new_dist
                    predecessors[v] = u
                    
        steps.append({
            'T': set(T),
            'selected_v': u,
            'distances': dict(distances),
            'highlight_v': u
        })
        
        if u == end:
            break
            
    # Tái thiết lập đường đi ngắn nhất
    path = []
    curr = end
    if distances[end] != float('inf'):
        while curr is not None:
            path.append(curr)
            curr = predecessors[curr]
        path.reverse()
        
    return steps, path, distances[end]

def generate_word_document(filename, vertices, steps, path, path_length, start_node, end_node):
    """Tạo tệp tài liệu Word chứa bảng lời giải Dijkstra tự động tô đỏ cột đỉnh được chọn."""
    doc = Document()
    
    # Thiết lập Font chữ mặc định là Times New Roman, cỡ 12pt (chuẩn học thuật)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Đoạn đề bài
    p_title = doc.add_paragraph()
    run_bold = p_title.add_run("Câu 3 (4 điểm) ")
    run_bold.bold = True
    p_title.add_run("Cho đồ thị liên thông, có trọng số. Dùng thuật toán Dijkstra để tìm đường đi ngắn nhất từ đỉnh ")
    p_title.add_run(start_node).bold = True
    p_title.add_run(" đến đỉnh ")
    p_title.add_run(end_node).bold = True
    p_title.add_run(":")
    
    # Lời giải mở đầu
    p_ans = doc.add_paragraph()
    run_ans_bold = p_ans.add_run("# Trả lời: ")
    run_ans_bold.bold = True
    p_ans.add_run("Trình bày cách giải bằng tay vào đây (yêu cầu trình bày dạng bảng):")
    
    # Thiết lập kích thước bảng: Số cột = 2 + số đỉnh, Số dòng = Số bước lưu vết + 1 dòng tiêu đề
    cols_count = 2 + len(vertices)
    rows_count = len(steps) + 1
    
    table = doc.add_table(rows=rows_count, cols=cols_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Đặt tiêu đề cột
    headers = ["T", "V_i"] + [f"D_{v}" for v in vertices]
    hdr_cells = table.rows[0].cells
    for col_idx, text in enumerate(headers):
        hdr_cells[col_idx].text = text
        p = hdr_cells[col_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        set_cell_borders(hdr_cells[col_idx])
        set_cell_margins(hdr_cells[col_idx])
        # Thiết lập chiều rộng cho ô tiêu đề
        if col_idx == 0:
            hdr_cells[col_idx].width = Cm(3.0)
        elif col_idx == 1:
            hdr_cells[col_idx].width = Cm(1.2)
        else:
            hdr_cells[col_idx].width = Cm(1.2)
        
    # Tập hợp các đỉnh đã duyệt qua
    selected_so_far = set()
    
    # Điền dữ liệu vào bảng
    for row_idx, step in enumerate(steps):
        row_cells = table.rows[row_idx + 1].cells
        
        # Thiết lập chiều rộng cho từng ô trong hàng
        row_cells[0].width = Cm(3.0)
        row_cells[1].width = Cm(1.2)
        for i in range(2, cols_count):
            row_cells[i].width = Cm(1.2)
            
        T_set = step['T']
        selected_v = step['selected_v']
        highlight_v = step['highlight_v']
        step_distances = step['distances']
        
        # Cột T
        if row_idx == 0:
            t_str = format_T(set(vertices), vertices)
        else:
            t_str = format_T(T_set, vertices)
        row_cells[0].text = t_str
        
        # Cột V_i
        vi_str = f"V_{selected_v}" if selected_v else "-"
        row_cells[1].text = vi_str
        
        # Căn giữa văn bản của cột T và V_i
        for i in (0, 1):
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_borders(row_cells[i])
            set_cell_margins(row_cells[i])
            
        # Điền khoảng cách của các đỉnh
        for col_idx, v in enumerate(vertices):
            cell = row_cells[2 + col_idx]
            
            if v == highlight_v:
                # Đỉnh này vừa được chọn để duyệt ở bước hiện tại
                dist_val = step_distances[v]
                cell_text = str(dist_val) if dist_val != float('inf') else "∞"
                cell.text = cell_text
                
                # Chuyển chữ màu đỏ in đậm, không tô màu nền (để ô màu trắng)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                selected_so_far.add(v)
            elif v in selected_so_far:
                # Đỉnh đã được duyệt từ các bước trước
                cell.text = "-"
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Đỉnh chưa được duyệt
                dist_val = step_distances[v]
                if dist_val == float('inf'):
                    cell.text = "∞"
                else:
                    cell.text = str(dist_val)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            set_cell_borders(cell)
            set_cell_margins(cell)
            
    # Ghi chú đường đi và độ dài đường đi ngắn nhất ở dưới bảng
    doc.add_paragraph() # Dòng trống ngăn cách
    
    p_path = doc.add_paragraph()
    p_path.add_run("# Trả lời: ").bold = True
    p_path.add_run("chỉ ra đường đi ngắn nhất: ")
    path_str = " <- ".join([f"V_{node}" for node in reversed(path)])
    p_path.add_run(f"P = {path_str}").bold = True
    
    p_len = doc.add_paragraph()
    p_len.add_run("# Trả lời: ").bold = True
    p_len.add_run("độ dài đường đi ngắn nhất ")
    p_len.add_run(f"D_{end_node} = {path_length}").bold = True
    
    doc.save(filename)
    print(f"\n[Aris Success] Đã xuất bảng kết quả Dijkstra thành công ra file: {filename}")

def main():
    print("=" * 70)
    print("      ARIS DIJKSTRA PATH FINDER & WORD EXPORTER TOOL 🎮")
    print("=" * 70)
    print("Sensei muốn chạy thử đồ thị mẫu mặc định trong ảnh không?")
    print("Đồ thị mẫu:")
    print("  Danh sách đỉnh: A, B, C, D, E, F, G")
    print("  Điểm bắt đầu: A -> Điểm đích: G")
    choice = input("Sensei chọn dùng đồ thị mặc định? (Y/n): ").strip().lower()
    
    if choice in ('', 'y', 'yes'):
        vertices = ['A', 'B', 'C', 'D', 'E', 'F', 'G']
        start_node = 'A'
        end_node = 'G'
        
        # Ma trận tam giác dưới mặc định tương ứng với hình ảnh
        matrix_input = [
            [0],
            [12, 0],
            [5, 3, 0],
            [float('inf'), 13, 2, 0],
            [float('inf'), 6, float('inf'), 5, 0],
            [float('inf'), float('inf'), 25, 15, 14, 0],
            [float('inf'), float('inf'), float('inf'), float('inf'), 7, 16, 0]
        ]
        
        adjacency_matrix = {v: {} for v in vertices}
        for u in vertices:
            for v in vertices:
                adjacency_matrix[u][v] = 0.0 if u == v else float('inf')
                
        for i in range(len(vertices)):
            for j in range(i + 1):
                val = matrix_input[i][j]
                u = vertices[i]
                v = vertices[j]
                adjacency_matrix[u][v] = val
                adjacency_matrix[v][u] = val
    else:
        print("\n--- Nhập Dữ Liệu Tự Chọn Cho Đồ Thị ---")
        vertices_str = input("Nhập danh sách đỉnh (cách nhau bằng dấu cách hoặc phẩy): ").strip()
        vertices = re.split(r'[\s,]+', vertices_str)
        vertices = [v.upper() for v in vertices if v]
        
        if not vertices:
            print("[Aris Error] Danh sách đỉnh không hợp lệ!")
            return
            
        print(f"Các đỉnh đã nhận diện: {', '.join(vertices)}")
        
        start_node = input(f"Đỉnh bắt đầu (mặc định: {vertices[0]}): ").strip().upper()
        if not start_node or start_node not in vertices:
            start_node = vertices[0]
            
        end_node = input(f"Đỉnh kết thúc (mặc định: {vertices[-1]}): ").strip().upper()
        if not end_node or end_node not in vertices:
            end_node = vertices[-1]
            
        print("\nNhập khoảng cách trực quan giữa các cặp đỉnh:")
        print("-> Nhập số thực hoặc số nguyên nếu có cạnh nối.")
        print("-> Chỉ nhấn ENTER (để trống) nếu không có cạnh nối.")
        
        adjacency_matrix = {v: {} for v in vertices}
        for u in vertices:
            for v in vertices:
                adjacency_matrix[u][v] = 0.0 if u == v else float('inf')
                
        # Duyệt qua từng cặp đỉnh u, v (u < v theo thứ tự danh sách)
        for i in range(len(vertices)):
            u = vertices[i]
            for j in range(i + 1, len(vertices)):
                v = vertices[j]
                while True:
                    val_str = input(f"Khoảng cách {u} - {v} (Enter nếu không nối): ").strip()
                    if not val_str:
                        val = float('inf')
                        break
                    try:
                        val = float(val_str)
                        if val.is_integer():
                            val = int(val)
                        break
                    except ValueError:
                        print("[Aris Error] Vui lòng nhập một số hợp lệ hoặc chỉ nhấn Enter!")
                
                adjacency_matrix[u][v] = val
                adjacency_matrix[v][u] = val

    print(f"\n[Dijkstra] Đang tìm đường đi ngắn nhất từ {start_node} đến {end_node}...")
    steps, path, path_length = run_dijkstra(vertices, adjacency_matrix, start_node, end_node)
    
    if not path:
        print(f"[Aris Error] Không có đường đi khả dụng từ {start_node} đến {end_node}!")
        return
        
    print(f"[Aris Info] Độ dài đường đi ngắn nhất: {path_length}")
    print(f"[Aris Info] Hành trình: {' -> '.join(path)}")
    
    output_filename = "dijkstra_result.docx"
    generate_word_document(output_filename, vertices, steps, path, path_length, start_node, end_node)
    print("=" * 70)
    print(f"🎉 Hoàn thành nhiệm vụ! File kết quả được lưu tại: {os.path.abspath(output_filename)}")
    print("=" * 70)

if __name__ == '__main__':
    main()
